"""Multi-format document parser for the RAG ingestion pipeline.

Supports TXT, Markdown, PDF (via pymupdf), DOCX (via python-docx), and HTML (via beautifulsoup4).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    import fitz as _fitz_module
except ImportError:
    _fitz_module = None  # type: ignore[assignment]

try:
    import docx as _docx_module
except ImportError:
    _docx_module = None  # type: ignore[assignment]

try:
    from bs4 import BeautifulSoup as _BeautifulSoup
except ImportError:
    _BeautifulSoup = None  # type: ignore[assignment]


@dataclass
class DocumentSection:
    heading: str
    level: int
    text: str
    page_start: int | None
    page_end: int | None
    section_id: str


@dataclass
class ParsedDocument:
    title: str
    text: str
    mime_type: str
    sections: list[DocumentSection]
    page_count: int | None
    metadata: dict[str, Any] = field(default_factory=dict)


SUPPORTED_MIME_TYPES: dict[str, str] = {
    "text/plain": "txt",
    "text/markdown": "md",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/html": "html",
}

_EXTENSION_TO_MIME: dict[str, str] = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".html": "text/html",
    ".htm": "text/html",
}


def _slugify(text: str) -> str:
    """Convert heading text to a URL-safe section identifier."""
    slug = text.lower().strip()
    slug = re.sub(r"[^a-z0-9\s-]", "", slug)
    slug = re.sub(r"[\s-]+", "-", slug)
    return slug.strip("-") or "section"


def _detect_mime_type(filename: str) -> str | None:
    ext = Path(filename).suffix.lower()
    return _EXTENSION_TO_MIME.get(ext)


def _parse_txt(file_bytes: bytes, filename: str) -> ParsedDocument:
    text = file_bytes.decode("utf-8", errors="replace")
    title = Path(filename).stem
    return ParsedDocument(
        title=title,
        text=text,
        mime_type="text/plain",
        sections=[],
        page_count=None,
    )


def _parse_md(file_bytes: bytes, filename: str) -> ParsedDocument:
    text = file_bytes.decode("utf-8", errors="replace")
    title = Path(filename).stem
    lines = text.splitlines()

    sections: list[DocumentSection] = []
    heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$")

    current_heading: str | None = None
    current_level: int = 0
    current_lines: list[str] = []

    def flush_section() -> None:
        if current_heading is not None:
            body = "\n".join(current_lines).strip()
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    level=current_level,
                    text=body,
                    page_start=None,
                    page_end=None,
                    section_id=_slugify(current_heading),
                )
            )

    for line in lines:
        m = heading_pattern.match(line)
        if m:
            flush_section()
            current_heading = m.group(2).strip()
            current_level = len(m.group(1))
            current_lines = []
        else:
            current_lines.append(line)

    flush_section()

    if sections and sections[0].level == 1:
        title = sections[0].heading

    return ParsedDocument(
        title=title,
        text=text,
        mime_type="text/markdown",
        sections=sections,
        page_count=None,
    )


def _parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    import io

    fitz = _fitz_module
    if fitz is None:
        raise ImportError("pymupdf is required for PDF parsing: pip install pymupdf")

    doc = fitz.open(stream=io.BytesIO(file_bytes), filetype="pdf")
    page_count = len(doc)
    full_text_parts: list[str] = []
    sections: list[DocumentSection] = []

    heading_re = re.compile(r"^[A-Z][A-Z0-9 \t:/-]{2,}$")

    current_heading: str | None = None
    current_page_start: int | None = None
    current_page_end: int | None = None
    current_lines: list[str] = []

    def flush_section(page_end: int) -> None:
        if current_heading is not None:
            body = "\n".join(current_lines).strip()
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    level=1,
                    text=body,
                    page_start=current_page_start,
                    page_end=page_end,
                    section_id=_slugify(current_heading),
                )
            )

    for page_num in range(page_count):
        page = doc[page_num]
        page_text = page.get_text()
        full_text_parts.append(page_text)

        for raw_line in page_text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if len(line) < 80 and heading_re.match(line):
                flush_section(page_num)
                current_heading = line
                current_page_start = page_num + 1
                current_page_end = page_num + 1
                current_lines = []
            else:
                current_lines.append(line)
                current_page_end = page_num + 1

    flush_section(current_page_end or page_count)

    full_text = "\n".join(full_text_parts)
    title = Path(filename).stem

    return ParsedDocument(
        title=title,
        text=full_text,
        mime_type="application/pdf",
        sections=sections,
        page_count=page_count,
    )


def _parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    import io

    python_docx = _docx_module
    if python_docx is None:
        raise ImportError(
            "python-docx is required for DOCX parsing: pip install python-docx"
        )

    doc = python_docx.Document(io.BytesIO(file_bytes))
    title = Path(filename).stem

    all_text_parts: list[str] = []
    sections: list[DocumentSection] = []

    current_heading: str | None = None
    current_level: int = 1
    current_lines: list[str] = []

    heading_style_re = re.compile(r"^Heading (\d+)$", re.IGNORECASE)

    def flush_section() -> None:
        if current_heading is not None:
            body = "\n".join(current_lines).strip()
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    level=current_level,
                    text=body,
                    page_start=None,
                    page_end=None,
                    section_id=_slugify(current_heading),
                )
            )

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        m = heading_style_re.match(style_name)
        if m:
            flush_section()
            current_heading = text
            current_level = int(m.group(1))
            current_lines = []
        else:
            current_lines.append(text)

        all_text_parts.append(text)

    flush_section()

    if sections and sections[0].level == 1:
        title = sections[0].heading

    full_text = "\n".join(all_text_parts)
    return ParsedDocument(
        title=title,
        text=full_text,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        sections=sections,
        page_count=None,
    )


def _parse_html(file_bytes: bytes, filename: str) -> ParsedDocument:
    BeautifulSoup = _BeautifulSoup
    if BeautifulSoup is None:
        raise ImportError(
            "beautifulsoup4 and lxml are required for HTML parsing: "
            "pip install beautifulsoup4 lxml"
        )

    html = file_bytes.decode("utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")

    title_tag = soup.find("title")
    title = title_tag.get_text(strip=True) if title_tag else Path(filename).stem

    body = soup.find("body") or soup
    full_text = body.get_text(separator="\n", strip=True)

    sections: list[DocumentSection] = []
    heading_tags = body.find_all(re.compile(r"^h[1-6]$"))

    for tag in heading_tags:
        level = int(tag.name[1])
        heading_text = tag.get_text(strip=True)

        sibling_parts: list[str] = []
        for sibling in tag.next_siblings:
            if hasattr(sibling, "name") and sibling.name and re.match(r"^h[1-6]$", sibling.name):
                break
            if hasattr(sibling, "get_text"):
                t = sibling.get_text(separator=" ", strip=True)
                if t:
                    sibling_parts.append(t)

        sections.append(
            DocumentSection(
                heading=heading_text,
                level=level,
                text=" ".join(sibling_parts).strip(),
                page_start=None,
                page_end=None,
                section_id=_slugify(heading_text),
            )
        )

    return ParsedDocument(
        title=title,
        text=full_text,
        mime_type="text/html",
        sections=sections,
        page_count=None,
    )


def parse_document(
    file_bytes: bytes, filename: str, mime_type: str | None = None
) -> ParsedDocument:
    """Parse a document from raw bytes into structured sections.

    Args:
        file_bytes: Raw file content.
        filename: Original filename used for MIME detection and title derivation.
        mime_type: Optional explicit MIME type; auto-detected from extension if not provided.

    Returns:
        ParsedDocument with extracted text, sections, and metadata.

    Raises:
        ValueError: If the file format is not supported.
    """
    if mime_type is None:
        mime_type = _detect_mime_type(filename)

    if mime_type is None or mime_type not in SUPPORTED_MIME_TYPES:
        raise ValueError(
            f"Unsupported file format: {filename!r} (mime_type={mime_type!r}). "
            f"Supported: {list(SUPPORTED_MIME_TYPES)}"
        )

    fmt = SUPPORTED_MIME_TYPES[mime_type]

    if fmt == "txt":
        return _parse_txt(file_bytes, filename)
    if fmt == "md":
        return _parse_md(file_bytes, filename)
    if fmt == "pdf":
        return _parse_pdf(file_bytes, filename)
    if fmt == "docx":
        return _parse_docx(file_bytes, filename)
    if fmt == "html":
        return _parse_html(file_bytes, filename)

    raise ValueError(f"Unhandled format: {fmt!r}")
